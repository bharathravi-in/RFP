"""
Export service for generating RFP response documents.
"""
import io
import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime


def setup_document_styles(doc):
    """
    Configure professional enterprise-ready document styles.
    Sets up heading styles, body text, and margins.
    """
    # Set document margins (1 inch = 914400 EMUs)
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1)
    
    # Configure Heading 1 style
    try:
        h1_style = doc.styles['Heading 1']
        h1_style.font.name = 'Calibri'
        h1_style.font.size = Pt(18)
        h1_style.font.bold = True
        h1_style.font.color.rgb = RGBColor(31, 73, 125)  # Professional blue
        h1_style.paragraph_format.space_before = Pt(24)
        h1_style.paragraph_format.space_after = Pt(12)
    except KeyError:
        pass
    
    # Configure Heading 2 style
    try:
        h2_style = doc.styles['Heading 2']
        h2_style.font.name = 'Calibri'
        h2_style.font.size = Pt(14)
        h2_style.font.bold = True
        h2_style.font.color.rgb = RGBColor(54, 95, 145)
        h2_style.paragraph_format.space_before = Pt(18)
        h2_style.paragraph_format.space_after = Pt(8)
    except KeyError:
        pass
    
    # Configure Normal paragraph style
    try:
        normal_style = doc.styles['Normal']
        normal_style.font.name = 'Calibri'
        normal_style.font.size = Pt(11)
        normal_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        normal_style.paragraph_format.line_spacing = 1.15
        normal_style.paragraph_format.space_after = Pt(8)
    except KeyError:
        pass


def safe_add_heading(doc, text, level=1):
    """
    Safely add a heading, falling back to manual styling if style doesn't exist.
    This handles templates that may lack standard Word heading styles.
    """
    try:
        return doc.add_heading(text, level)
    except KeyError:
        # Fallback: create paragraph with manual heading styling
        para = doc.add_paragraph()
        run = para.add_run(text)
        
        if level == 0:  # Title
            run.font.size = Pt(28)
            run.font.bold = True
            run.font.color.rgb = RGBColor(31, 73, 125)
        elif level == 1:
            run.font.size = Pt(18)
            run.font.bold = True
            run.font.color.rgb = RGBColor(31, 73, 125)
        elif level == 2:
            run.font.size = Pt(14)
            run.font.bold = True
            run.font.color.rgb = RGBColor(54, 95, 145)
        else:
            run.font.size = Pt(12)
            run.font.bold = True
        
        run.font.name = 'Calibri'
        return para


def safe_add_paragraph(doc, text='', style=None):
    """
    Safely add a paragraph with optional style, falling back to manual styling if style doesn't exist.
    This handles templates that may lack standard Word paragraph styles.
    """
    try:
        if style:
            para = doc.add_paragraph(text, style=style)
        else:
            para = doc.add_paragraph(text)
        return para
    except KeyError:
        # Fallback: create paragraph without style, apply manual formatting
        para = doc.add_paragraph()
        if text:
            run = para.add_run(text)
            run.font.name = 'Calibri'
            run.font.size = Pt(11)
            # Add bullet character for list styles
            if style and 'bullet' in style.lower():
                run.text = '• ' + text
            elif style and 'number' in style.lower():
                run.text = text
        return para


def add_document_header(doc, project_name, organization_name=None):
    """
    Add professional header with company name and project to all pages.
    """
    for section in doc.sections:
        header = section.header
        header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        
        # Left side: Company/Organization name
        if organization_name:
            run = header_para.add_run(organization_name)
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(100, 100, 100)
            run.font.bold = True
            
            run = header_para.add_run('    |    ')
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(180, 180, 180)
        
        # Right side: Project name
        run = header_para.add_run(project_name[:40] + ('...' if len(project_name) > 40 else ''))
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(128, 128, 128)
        run.italic = True
        
        header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_section_divider(doc):
    """Add a professional section divider line."""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run('─' * 40)
    run.font.color.rgb = RGBColor(200, 200, 200)
    run.font.size = Pt(10)


def style_table(table):
    """Apply professional styling to a table."""
    # Set table alignment
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Style header row
    if table.rows:
        header_row = table.rows[0]
        for cell in header_row.cells:
            # Set background color for header
            shading = OxmlElement('w:shd')
            shading.set(qn('w:fill'), '1F497D')  # Dark blue
            cell._tc.get_or_add_tcPr().append(shading)
            
            # Set text color to white
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.color.rgb = RGBColor(255, 255, 255)
                    run.font.bold = True
    
    # Style alternating rows
    for idx, row in enumerate(table.rows[1:], 1):
        if idx % 2 == 0:
            for cell in row.cells:
                shading = OxmlElement('w:shd')
                shading.set(qn('w:fill'), 'F5F5F5')  # Light gray
                cell._tc.get_or_add_tcPr().append(shading)


def add_markdown_to_doc(doc, content):
    """
    Convert markdown content to Word document formatting.
    Handles: headers, bold, italic, lists, code blocks, and paragraphs.
    """
    if not content:
        return
    
    lines = content.split('\n')
    current_list_items = []
    is_in_list = False
    is_in_code_block = False
    code_block_content = []
    code_block_language = ''
    
    def flush_list():
        nonlocal current_list_items, is_in_list
        if current_list_items:
            for item in current_list_items:
                # Remove list markers
                clean_item = re.sub(r'^[\*\-]\s*|^\d+\.\s*', '', item)
                para = safe_add_paragraph(doc, style='List Bullet')
                add_formatted_text(para, clean_item)
            current_list_items = []
            is_in_list = False
    
    def add_formatted_text(paragraph, text):
        """Add text with inline formatting (bold, italic)."""
        # Pattern for bold, italic, code
        pattern = r'(\*\*\*(.+?)\*\*\*|\*\*(.+?)\*\*|\*(.+?)\*|`(.+?)`|\[(.+?)\]\((.+?)\))'
        last_end = 0
        
        for match in re.finditer(pattern, text):
            # Add text before match
            if match.start() > last_end:
                paragraph.add_run(text[last_end:match.start()])
            
            if match.group(2):  # Bold italic ***text***
                run = paragraph.add_run(match.group(2))
                run.bold = True
                run.italic = True
            elif match.group(3):  # Bold **text**
                run = paragraph.add_run(match.group(3))
                run.bold = True
            elif match.group(4):  # Italic *text*
                run = paragraph.add_run(match.group(4))
                run.italic = True
            elif match.group(5):  # Code `text`
                run = paragraph.add_run(match.group(5))
                run.font.name = 'Courier New'
                run.font.size = Pt(10)
            elif match.group(6) and match.group(7):  # Link [text](url)
                run = paragraph.add_run(match.group(6))
                run.underline = True
            
            last_end = match.end()
        
        # Add remaining text
        if last_end < len(text):
            paragraph.add_run(text[last_end:])
    
    for line in lines:
        stripped = line.strip()
        
        # Handle code blocks (including mermaid) - check both original and stripped line
        if stripped.startswith('```') or line.lstrip().startswith('```'):
            if is_in_code_block:
                # End of code block
                flush_list()
                if code_block_language == 'mermaid':
                    # For mermaid, render the diagram as an image
                    mermaid_code = '\n'.join(code_block_content)
                    try:
                        from .mermaid_service import render_mermaid_to_bytes_io
                        diagram_buffer = render_mermaid_to_bytes_io(mermaid_code)
                        if diagram_buffer:
                            # Add the diagram image to the document
                            doc.add_picture(diagram_buffer, width=Inches(5.5))
                            # Add caption
                            caption_para = doc.add_paragraph()
                            caption_run = caption_para.add_run('Figure: Architecture Diagram')
                            caption_run.italic = True
                            caption_run.font.size = Pt(10)
                            caption_run.font.color.rgb = RGBColor(100, 100, 100)
                            caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        else:
                            # Fallback: Add placeholder text if rendering failed
                            para = doc.add_paragraph()
                            run = para.add_run('[Diagram could not be rendered - see web application]')
                            run.italic = True
                            run.font.size = Pt(10)
                            run.font.color.rgb = RGBColor(100, 100, 100)
                    except Exception as e:
                        # Error handling: Add placeholder text
                        para = doc.add_paragraph()
                        run = para.add_run(f'[Diagram rendering error - see web application]')
                        run.italic = True
                        run.font.size = Pt(10)
                        run.font.color.rgb = RGBColor(100, 100, 100)
                elif code_block_content:
                    # For other non-empty code blocks, add as monospace text
                    para = doc.add_paragraph()
                    code_text = '\n'.join(code_block_content)
                    run = para.add_run(code_text)
                    run.font.name = 'Courier New'
                    run.font.size = Pt(9)
                code_block_content = []
                is_in_code_block = False
                code_block_language = ''
            else:
                # Start of code block - extract language
                flush_list()
                is_in_code_block = True
                # Get the part after ```
                lang_part = stripped[3:] if stripped.startswith('```') else line.lstrip()[3:]
                code_block_language = lang_part.strip().lower()
            continue
        
        if is_in_code_block:
            code_block_content.append(line)
            continue
        
        # Skip empty lines
        if not stripped:
            flush_list()
            continue
        
        # Headers
        if stripped.startswith('### '):
            flush_list()
            heading = doc.add_heading(level=3)
            add_formatted_text(heading, stripped[4:])
            continue
        elif stripped.startswith('## '):
            flush_list()
            heading = doc.add_heading(level=2)
            add_formatted_text(heading, stripped[3:])
            continue
        elif stripped.startswith('# '):
            flush_list()
            heading = doc.add_heading(level=1)
            add_formatted_text(heading, stripped[2:])
            continue
        
        # Lists (bullet or numbered)
        if re.match(r'^[\*\-]\s', stripped) or re.match(r'^\d+\.\s', stripped):
            is_in_list = True
            current_list_items.append(stripped)
            continue
        
        # Horizontal rule
        if re.match(r'^[-*_]{3,}$', stripped):
            flush_list()
            doc.add_paragraph('_' * 50)
            continue
        
        # Regular paragraph
        flush_list()
        para = doc.add_paragraph()
        add_formatted_text(para, stripped)
    
    # Flush any remaining list
    flush_list()


def generate_docx(project, questions):
    """
    Generate a DOCX document with project responses.
    
    Args:
        project: Project model instance
        questions: List of Question model instances with answers
    
    Returns:
        BytesIO buffer containing the DOCX file
    """
    doc = Document()
    
    # Title
    title = doc.add_heading(f'RFP Response: {project.name}', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Metadata
    doc.add_paragraph(f'Generated: {datetime.now().strftime("%B %d, %Y")}')
    doc.add_paragraph(f'Total Questions: {len(questions)}')
    approved_count = sum(1 for q in questions if q.status == 'approved')
    doc.add_paragraph(f'Approved Answers: {approved_count}')
    
    doc.add_paragraph()  # Space
    
    # Group questions by section
    sections = {}
    for q in questions:
        section = q.section or 'General'
        if section not in sections:
            sections[section] = []
        sections[section].append(q)
    
    # Add each section
    for section_name, section_questions in sections.items():
        doc.add_heading(section_name, level=1)
        
        for i, question in enumerate(section_questions, 1):
            # Question
            q_para = doc.add_paragraph()
            q_run = q_para.add_run(f'Q{i}: {question.text}')
            q_run.bold = True
            
            # Answer
            answer = question.current_answer
            if answer and answer.content:
                add_markdown_to_doc(doc, answer.content)
                
                # Status indicator
                status_para = doc.add_paragraph()
                status_run = status_para.add_run(f'Status: {question.status.upper()}')
                status_run.italic = True
                status_run.font.size = Pt(10)
                
                if answer.confidence_score:
                    conf_run = status_para.add_run(f' | Confidence: {int(answer.confidence_score * 100)}%')
                    conf_run.italic = True
                    conf_run.font.size = Pt(10)
            else:
                no_answer = doc.add_paragraph('No answer provided.')
                no_answer.runs[0].italic = True
            
            doc.add_paragraph()  # Space between questions
    
    # Save to buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    return buffer


def generate_xlsx(project, questions):
    """
    Generate an XLSX spreadsheet with project responses.
    """
    import pandas as pd
    
    data = []
    for i, q in enumerate(questions, 1):
        answer = q.current_answer
        data.append({
            'No.': i,
            'Section': q.section or 'General',
            'Question': q.text,
            'Answer': answer.content if answer else '',
            'Status': q.status,
            'Confidence': f"{int(answer.confidence_score * 100)}%" if answer and answer.confidence_score else '',
            'AI Generated': 'Yes' if answer and answer.is_ai_generated else 'No'
        })
    
    df = pd.DataFrame(data)
    
    buffer = io.BytesIO()
    with pd.ExcelWriter(buffer, engine='openpyxl') as writer:
        df.to_excel(writer, sheet_name='RFP Responses', index=False)
    
    buffer.seek(0)
    return buffer


def add_page_footer(doc, version_text):
    """Add page footer with page numbers and version to all sections."""
    for section in doc.sections:
        footer = section.footer
        footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add version on left
        run = footer_para.add_run(f'{version_text}')
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(128, 128, 128)
        
        # Add separator
        run = footer_para.add_run('    |    ')
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(128, 128, 128)
        
        # Add page number field
        run = footer_para.add_run('Page ')
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(128, 128, 128)
        
        # PAGE field
        fld_char1 = OxmlElement('w:fldChar')
        fld_char1.set(qn('w:fldCharType'), 'begin')
        instr_text = OxmlElement('w:instrText')
        instr_text.text = 'PAGE'
        fld_char2 = OxmlElement('w:fldChar')
        fld_char2.set(qn('w:fldCharType'), 'separate')
        fld_char3 = OxmlElement('w:fldChar')
        fld_char3.set(qn('w:fldCharType'), 'end')
        
        run2 = footer_para.add_run()
        run2._r.append(fld_char1)
        run2._r.append(instr_text)
        run2._r.append(fld_char2)
        run2._r.append(fld_char3)
        
        run = footer_para.add_run(' of ')
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(128, 128, 128)
        
        # NUMPAGES field
        fld_char1 = OxmlElement('w:fldChar')
        fld_char1.set(qn('w:fldCharType'), 'begin')
        instr_text = OxmlElement('w:instrText')
        instr_text.text = 'NUMPAGES'
        fld_char2 = OxmlElement('w:fldChar')
        fld_char2.set(qn('w:fldCharType'), 'separate')
        fld_char3 = OxmlElement('w:fldChar')
        fld_char3.set(qn('w:fldCharType'), 'end')
        
        run3 = footer_para.add_run()
        run3._r.append(fld_char1)
        run3._r.append(instr_text)
        run3._r.append(fld_char2)
        run3._r.append(fld_char3)


def add_vendor_visibility_section(doc, project, organization=None):
    """
    Add Vendor Visibility / Eligibility section to the proposal.
    Surfaces vendor credentials and eligibility criteria.
    """
    doc.add_heading('Vendor Profile & Eligibility', level=1)
    
    doc.add_paragraph()
    
    # Introduction text
    intro = doc.add_paragraph()
    intro_run = intro.add_run('This section provides an overview of our organization\'s credentials and eligibility to fulfill the requirements of this RFP.')
    intro_run.font.size = Pt(10)
    intro_run.italic = True
    intro_run.font.color.rgb = RGBColor(100, 100, 100)
    
    doc.add_paragraph()
    
    # Get vendor profile data from organization settings or defaults
    org_settings = {}
    if organization and hasattr(organization, 'settings') and organization.settings:
        org_settings = organization.settings
    
    vendor_profile = org_settings.get('vendor_profile', {})
    
    # Company Information Table
    doc.add_heading('Company Information', level=2)
    
    company_info = [
        ('Company Name', vendor_profile.get('company_name', organization.name if organization else 'Not specified')),
        ('Registration Country', vendor_profile.get('registration_country', 'Not specified')),
        ('Years in Business', vendor_profile.get('years_in_business', 'Not specified')),
        ('Employee Count', vendor_profile.get('employee_count', 'Not specified')),
        ('Headquarters', vendor_profile.get('headquarters', 'Not specified')),
    ]
    
    # Create table for company info
    table = doc.add_table(rows=len(company_info), cols=2)
    table.style = 'Table Grid'
    
    for i, (label, value) in enumerate(company_info):
        row = table.rows[i]
        cell1 = row.cells[0]
        cell2 = row.cells[1]
        
        run1 = cell1.paragraphs[0].add_run(label)
        run1.bold = True
        run1.font.size = Pt(10)
        
        run2 = cell2.paragraphs[0].add_run(str(value))
        run2.font.size = Pt(10)
    
    doc.add_paragraph()
    
    # Certifications & Compliance
    doc.add_heading('Certifications & Compliance', level=2)
    
    certifications = vendor_profile.get('certifications', [])
    if not certifications:
        # Default certifications if not specified
        certifications = ['SOC 2 Type II', 'ISO 27001', 'GDPR Compliant']
    
    for cert in certifications:
        cert_para = safe_add_paragraph(doc, style='List Bullet')
        cert_run = cert_para.add_run(f'✓ {cert}')
        cert_run.font.size = Pt(10)
    
    doc.add_paragraph()
    
    # Industry Experience
    doc.add_heading('Industry Experience', level=2)
    
    industries = vendor_profile.get('industries', [])
    if not industries:
        # Check if project has industry info
        if hasattr(project, 'industry') and project.industry:
            industries = [project.industry]
        else:
            industries = ['Technology', 'Financial Services', 'Healthcare']
    
    industries_para = doc.add_paragraph()
    industries_run = industries_para.add_run('Our organization has extensive experience serving clients in the following industries:')
    industries_run.font.size = Pt(10)
    
    for industry in industries:
        ind_para = safe_add_paragraph(doc, style='List Bullet')
        ind_run = ind_para.add_run(industry)
        ind_run.font.size = Pt(10)
    
    doc.add_paragraph()
    
    # Geographic Presence
    doc.add_heading('Geographic Presence', level=2)
    
    geographies = vendor_profile.get('geographies', [])
    if not geographies:
        # Check if project has geography info
        if hasattr(project, 'geography') and project.geography:
            geographies = [project.geography]
        else:
            geographies = ['Global', 'North America', 'Europe', 'Asia Pacific']
    
    geo_para = doc.add_paragraph()
    geo_run = geo_para.add_run('We have operational presence and capability to deliver services in:')
    geo_run.font.size = Pt(10)
    
    for geo in geographies:
        geo_item = safe_add_paragraph(doc, style='List Bullet')
        geo_item_run = geo_item.add_run(geo)
        geo_item_run.font.size = Pt(10)
    
    doc.add_page_break()


def generate_proposal_docx(project, sections, include_qa=True, questions=None, organization=None, template_path=None):
    """
    Generate a full proposal DOCX with all sections.
    
    Args:
        project: Project model instance
        sections: List of RFPSection model instances
        include_qa: Whether to include Q&A section
        questions: List of Question model instances (if include_qa is True)
        organization: Organization model instance for vendor profile
        template_path: Optional path to DOCX template file to use as base
    
    Returns:
        BytesIO buffer containing the DOCX file
    """
    import os
    import logging
    logger = logging.getLogger(__name__)
    
    # Track if we're using a template (affects style handling)
    using_template = False
    
    # Load template or create new document
    if template_path and os.path.exists(template_path):
        try:
            doc = Document(template_path)
            logger.info(f"Using DOCX template: {template_path}")
            using_template = True
            # Clear existing content from template but keep styles and section properties
            # We must preserve sectPr elements for table width calculations
            from docx.oxml.ns import qn
            for element in doc.element.body[:]:
                # Don't remove sectPr (section properties) - needed for page layout/table widths
                if element.tag != qn('w:sectPr'):
                    doc.element.body.remove(element)
            # Ensure required styles exist (some templates may lack them)
            try:
                setup_document_styles(doc)
            except Exception as style_err:
                logger.warning(f"Could not setup styles on template: {style_err}")
        except Exception as e:
            logger.warning(f"Failed to load template {template_path}: {e}, using blank document")
            doc = Document()
            setup_document_styles(doc)
    else:
        doc = Document()
        # ========================================
        # ENTERPRISE FORMATTING SETUP
        # ========================================
        # Apply professional styles and margins
        setup_document_styles(doc)
    
    # Get organization name for headers
    org_name = None
    if organization and hasattr(organization, 'name'):
        org_name = organization.name
    
    # Get version info
    current_date = datetime.now()
    version_number = getattr(project, 'version', 1) or 1
    version_text = f"Version {version_number}.0 – {current_date.strftime('%B %Y')}"
    
    # ========================================
    # TITLE PAGE
    # ========================================
    
    # Add some spacing at top
    for _ in range(3):
        doc.add_paragraph()
    
    # Draft label
    draft_label = doc.add_paragraph()
    draft_run = draft_label.add_run('DRAFT PROPOSAL')
    draft_run.font.size = Pt(14)
    draft_run.font.color.rgb = RGBColor(128, 128, 128)
    draft_run.bold = True
    draft_label.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    # Main Title
    title = safe_add_heading(doc, project.name, 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Subtitle
    subtitle = doc.add_paragraph()
    subtitle_run = subtitle.add_run('Proposal Response')
    subtitle_run.font.size = Pt(20)
    subtitle_run.font.color.rgb = RGBColor(75, 0, 130)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Client info section
    if hasattr(project, 'client_name') and project.client_name:
        prepared_for = doc.add_paragraph()
        prepared_for_run = prepared_for.add_run('Prepared For:')
        prepared_for_run.font.size = Pt(12)
        prepared_for_run.bold = True
        prepared_for.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        client_para = doc.add_paragraph()
        client_run = client_para.add_run(project.client_name)
        client_run.font.size = Pt(14)
        client_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()
    
    # Version and Date
    version_para = doc.add_paragraph()
    version_run = version_para.add_run(version_text)
    version_run.font.size = Pt(12)
    version_run.font.color.rgb = RGBColor(100, 100, 100)
    version_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Generated date
    date_para = doc.add_paragraph()
    date_run = date_para.add_run(f'Generated: {current_date.strftime("%B %d, %Y at %I:%M %p")}')
    date_run.font.size = Pt(10)
    date_run.font.color.rgb = RGBColor(128, 128, 128)
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Confidential notice at bottom
    for _ in range(5):
        doc.add_paragraph()
    
    confidential = doc.add_paragraph()
    conf_run = confidential.add_run('CONFIDENTIAL')
    conf_run.font.size = Pt(10)
    conf_run.font.color.rgb = RGBColor(192, 0, 0)
    conf_run.bold = True
    confidential.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    conf_note = doc.add_paragraph()
    conf_note_run = conf_note.add_run('This proposal contains confidential information intended only for the named recipient.')
    conf_note_run.font.size = Pt(9)
    conf_note_run.italic = True
    conf_note_run.font.color.rgb = RGBColor(128, 128, 128)
    conf_note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Page break after title
    doc.add_page_break()
    
    # ========================================
    # TABLE OF CONTENTS
    # ========================================
    
    toc_heading = doc.add_heading('Table of Contents', level=1)
    toc_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    doc.add_paragraph()
    
    # Add Vendor Profile to TOC first
    toc_item = doc.add_paragraph()
    toc_run = toc_item.add_run('Vendor Profile & Eligibility')
    toc_run.font.size = Pt(11)
    toc_run.bold = True
    toc_run2 = toc_item.add_run('  ' + '.' * 60)
    toc_run2.font.size = Pt(11)
    toc_run2.font.color.rgb = RGBColor(200, 200, 200)
    
    doc.add_paragraph()
    
    # Count sections with content
    sections_with_content = [(i, s) for i, s in enumerate(sections, 1) if s.content]
    
    for idx, (num, section) in enumerate(sections_with_content):
        toc_item = doc.add_paragraph()
        toc_item.paragraph_format.tab_stops.add_tab_stop(Inches(5.5))
        
        # Section number and title
        section_num = f'{idx + 1}.0'
        toc_run = toc_item.add_run(f'{section_num}  {section.title}')
        toc_run.font.size = Pt(11)
        
        # Add dotted leader (simulated with periods)
        toc_run2 = toc_item.add_run('  ' + '.' * 60)
        toc_run2.font.size = Pt(11)
        toc_run2.font.color.rgb = RGBColor(200, 200, 200)
    
    # Add Q&A to TOC
    if include_qa and questions:
        doc.add_paragraph()
        toc_item = doc.add_paragraph()
        qa_num = len(sections_with_content) + 1
        toc_run = toc_item.add_run(f'{qa_num}.0  Questions & Answers')
        toc_run.font.size = Pt(11)
        
        toc_run2 = toc_item.add_run('  ' + '.' * 60)
        toc_run2.font.size = Pt(11)
        toc_run2.font.color.rgb = RGBColor(200, 200, 200)
    
    doc.add_page_break()
    
    # ========================================
    # VENDOR VISIBILITY SECTION
    # ========================================
    
    add_vendor_visibility_section(doc, project, organization)
    
    # ========================================
    # PROPOSAL SECTIONS
    # ========================================
    
    for idx, (num, section) in enumerate(sections_with_content):
        # Section number and heading
        section_num = f'{idx + 1}.0'
        section_icon = section.section_type.icon if section.section_type else ''
        section_title = f'{section_num}  {section_icon} {section.title}'.strip()
        
        heading = doc.add_heading(section_title, level=1)
        
        # Section content - convert markdown to Word formatting
        add_markdown_to_doc(doc, section.content)
        
        # Section metadata footer
        if section.confidence_score:
            doc.add_paragraph()
            meta = doc.add_paragraph()
            meta_run = meta.add_run(f'AI Confidence: {int(section.confidence_score * 100)}%')
            meta_run.font.size = Pt(9)
            meta_run.italic = True
            meta_run.font.color.rgb = RGBColor(128, 128, 128)
        
        doc.add_page_break()
    
    # ========================================
    # Q&A SECTION (Optional)
    # ========================================
    
    if include_qa and questions:
        qa_num = len(sections_with_content) + 1
        doc.add_heading(f'{qa_num}.0  Questions & Answers', level=1)
        
        # Group by section
        qa_sections = {}
        for q in questions:
            section_name = q.section or 'General'
            if section_name not in qa_sections:
                qa_sections[section_name] = []
            qa_sections[section_name].append(q)
        
        # Summary stats
        answered_count = sum(1 for q in questions if q.current_answer and q.current_answer.content)
        summary = doc.add_paragraph()
        summary_run = summary.add_run(f'Total Questions: {len(questions)} | Answered: {answered_count}')
        summary_run.font.size = Pt(10)
        summary_run.italic = True
        summary_run.font.color.rgb = RGBColor(100, 100, 100)
        
        doc.add_paragraph()
        
        for section_name, section_questions in qa_sections.items():
            if len(qa_sections) > 1:
                doc.add_heading(section_name, level=2)
            
            for i, question in enumerate(section_questions, 1):
                # Question
                q_para = doc.add_paragraph()
                q_run = q_para.add_run(f'Q{i}: {question.text}')
                q_run.bold = True
                
                # Answer
                answer = question.current_answer
                if answer and answer.content:
                    add_markdown_to_doc(doc, answer.content)
                else:
                    no_answer = doc.add_paragraph('No answer provided.')
                    no_answer.runs[0].italic = True
                    no_answer.runs[0].font.color.rgb = RGBColor(150, 150, 150)
                
                doc.add_paragraph()
    
    # ========================================
    # ADD PAGE HEADERS AND FOOTERS
    # ========================================
    
    # Add professional header with company name and project
    add_document_header(doc, project.name, org_name)
    
    # Add page footer with version and page numbers
    add_page_footer(doc, version_text)
    
    # Save to buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    return buffer


def generate_proposal_xlsx(project, sections, questions=None):
    """
    Generate proposal XLSX with sections and Q&A.
    """
    import pandas as pd
    
    buffer = io.BytesIO()
    
    with pd.ExcelWriter(buffer, engine='openpyxl') as writer:
        # Sections sheet
        sections_data = []
        for i, section in enumerate(sections, 1):
            sections_data.append({
                'Order': i,
                'Section Type': section.section_type.name if section.section_type else 'Custom',
                'Title': section.title,
                'Content': section.content or '',
                'Status': section.status,
                'Confidence': f"{int(section.confidence_score * 100)}%" if section.confidence_score else '',
                'Version': section.version,
            })
        
        if sections_data:
            df_sections = pd.DataFrame(sections_data)
            df_sections.to_excel(writer, sheet_name='Proposal Sections', index=False)
        
        # Q&A sheet
        if questions:
            qa_data = []
            for i, q in enumerate(questions, 1):
                answer = q.current_answer
                qa_data.append({
                    'No.': i,
                    'Section': q.section or 'General',
                    'Question': q.text,
                    'Answer': answer.content if answer else '',
                    'Status': q.status,
                    'Confidence': f"{int(answer.confidence_score * 100)}%" if answer and answer.confidence_score else '',
                })
            
            if qa_data:
                df_qa = pd.DataFrame(qa_data)
                df_qa.to_excel(writer, sheet_name='Q&A Responses', index=False)
    
    buffer.seek(0)
    return buffer
